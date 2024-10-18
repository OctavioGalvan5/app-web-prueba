import os                 # Para manipular rutas y archivos
import tempfile           # Para crear archivos temporales
from datetime import datetime, timedelta  # Para manejar fechas y horas
from docxtpl import DocxTemplate  # Para trabajar con plantillas de Word
from docx import Document  # Para manipular documentos de Word
from flask import send_file  # Para enviar archivos en respuestas HTTP
from docx.shared import Inches
from models.calculos import calcular_porcentajes, formatear_dinero, transformar_fecha
from models.database import obtener_acordada, obtener_valor_uma
class Documento:
        def __init__(self, autos, expediente, periodo_desde, periodo_hasta,
                     fecha_de_cierre_de_liquidacion, fecha_de_regulacion,
                     fecha_aprobacion_sentencia, monto_aprobado,
                     monto_aprobado_actualizado, deuda, intereses,
                     imagenCapturaSentencia, imagenMonto):
            self.autos = autos
            self.expediente = expediente
            self.periodo_desde = periodo_desde
            self.periodo_hasta = periodo_hasta
            self.fecha_de_cierre_de_liquidacion = fecha_de_cierre_de_liquidacion
            self.fecha_de_regulacion = fecha_de_regulacion
            self.fecha_aprobacion_sentencia = fecha_aprobacion_sentencia
            self.monto_aprobado = monto_aprobado
            self.monto_aprobado_actualizado = monto_aprobado_actualizado
            self.deuda = deuda
            self.intereses = intereses
            self.imagenCapturaSentencia = imagenCapturaSentencia
            self.imagenMonto = imagenMonto

        def obtener_datos(self):
              datos = {}
              datos['Acordada_fecha_de_cierre_de_liquidacion'] = obtener_acordada(self.fecha_de_cierre_de_liquidacion)
              datos['UMA_fecha_de_cierre_de_liquidacion'] = obtener_valor_uma(self.fecha_de_cierre_de_liquidacion)
              datos['porcentajesFCL'], datos['cantidadFCL'], datos['minimoFCL'], datos['apoderadoFCL'], datos['reduccionFCL'], datos['ejecucionFCL'], datos['incidenciaFCL'] = calcular_porcentajes(self.monto_aprobado, datos['UMA_fecha_de_cierre_de_liquidacion'])

              datos['Acordada_fecha_de_regulacion'] = obtener_acordada(self.fecha_de_regulacion)
              datos['UMA_fecha_de_regulacion'] = obtener_valor_uma(self.fecha_de_regulacion)
              datos['porcentajesR'], datos['cantidadR'], datos['minimoR'], datos['apoderadoR'], datos['reduccionR'], datos['ejecucionR'], datos['incidenciaR'] = calcular_porcentajes(self.monto_aprobado, datos['UMA_fecha_de_regulacion'])
              

              datos['Acordada_fecha_aprobacion_sentencia'] = obtener_acordada(self.fecha_aprobacion_sentencia)
              datos['UMA_fecha_aprobacion_sentencia'] = obtener_valor_uma(self.fecha_aprobacion_sentencia)
              datos['porcentajesAS'], datos['cantidadAS'], datos['minimoAS'], datos['apoderadoAS'], datos['reduccionAS'], datos['ejecucionAS'], datos['incidenciaAS'] = calcular_porcentajes(self.monto_aprobado, datos['UMA_fecha_aprobacion_sentencia'])
              

              datos['porcentajesTP'], datos['cantidadTP'], datos['minimoTP'], datos['apoderadoTP'], datos['reduccionTP'], datos['ejecucionTP'], datos['incidenciaTP'] = calcular_porcentajes(self.monto_aprobado_actualizado, datos['UMA_fecha_de_regulacion'])

              datos['autos'] = self.autos
              datos['expediente'] = self.expediente
              datos['periodo_desde'] = transformar_fecha(self.periodo_desde)
              datos['periodo_hasta'] = transformar_fecha(self.periodo_hasta)
              datos['fecha_de_cierre_de_liquidacion'] = transformar_fecha(self.fecha_de_cierre_de_liquidacion)
              datos['fecha_de_regulacion'] = transformar_fecha(self.fecha_de_regulacion)
              datos['fecha_aprobacion_sentencia'] = transformar_fecha(self.fecha_aprobacion_sentencia)
              datos['monto_aprobado'] = formatear_dinero(float(self.monto_aprobado))
              datos['monto_aprobado_actualizado'] = formatear_dinero(float(self.monto_aprobado_actualizado))

              datos['deuda'] = formatear_dinero(float(self.deuda))
              datos['intereses'] = formatear_dinero(float(self.intereses))
              datos['imagenCapturaSentencia'] = self.imagenCapturaSentencia
              datos['imagenMonto'] = self.imagenMonto

              datos['UMA_fecha_de_regulacion'] = formatear_dinero(obtener_valor_uma(self.fecha_de_regulacion))
              datos['UMA_fecha_de_cierre_de_liquidacion'] = formatear_dinero(obtener_valor_uma(self.fecha_de_cierre_de_liquidacion))
              datos['UMA_fecha_aprobacion_sentencia'] = formatear_dinero(obtener_valor_uma(self.fecha_aprobacion_sentencia))

            
              return datos

        def crear_archivo_word(self, imagenes_por_marcador):
          
            plantilla = 'datos/calculadora_uma/plantilla_escrito_uma.docx'

            # Cargar la plantilla y crear el contexto
            doc = DocxTemplate(plantilla)
            contexto = self.obtener_datos()

            # Renderizar el documento con el contexto
            doc.render(contexto)

            # Guardar el documento editado temporalmente
            temp_doc_path = 'datos/calculadora_uma/documento_temporal.docx'
            doc.save(temp_doc_path)

            # Manejo de las imágenes
            doc = Document(temp_doc_path)

            # Insertar imágenes en posiciones específicas según el diccionario
            for marcador, imagen_path in imagenes_por_marcador.items():
                for paragraph in doc.paragraphs:
                    if marcador in paragraph.text:  # Buscar el marcador
                        paragraph.text = paragraph.text.replace(marcador, '')  # Eliminar el marcador del texto

                        if imagen_path:  # Verificar si hay imagen para ese marcador
                            run = paragraph.add_run()
                            run.add_picture(imagen_path, width=Inches(5))  # Ajusta el tamaño según sea necesario
                        break  # Detener la búsqueda en los párrafos una vez encontrado el marcador

            # Guardar el documento final editado
            final_path = 'datos/calculadora_uma/documento_editado.docx'
            doc.save(final_path)

            # Devolver la ruta del archivo final
            return final_path

        def procesar_imagenes(self):
            """Procesa las imágenes subidas por el usuario y las asigna a sus respectivos marcadores."""
            # Obtener las imágenes proporcionadas por el usuario (pueden ser opcionales)
            imagencaptura = self.imagenCapturaSentencia
            imagenmonto = self.imagenMonto
          

            # Crear diccionario para asignar las imágenes a sus marcadores
            imagenes_por_marcador = {}

            # Asignar cada imagen a su respectivo marcador si existe
            for imagen, marcador in zip(
                [imagencaptura, imagenmonto],
                ['Imagen_captura_aqui', 'Imagen_monto_aqui']
            ):
                if imagen:  # Verificar si el usuario subió una imagen
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(imagen.read())
                        imagenes_por_marcador[marcador] = temp_file.name  # Mapea el marcador a la ruta del archivo temporal
                else:
                    imagenes_por_marcador[marcador] = None  # No hay imagen para este marcador, se eliminará del documento

            # Llama a la función para crear el documento Word
            documento_path = self.crear_archivo_word(imagenes_por_marcador)  # Pasa el diccionario

            # Eliminar los archivos temporales después de usarlos
            for temp_file_path in imagenes_por_marcador.values():
                if temp_file_path:  # Solo eliminar archivos que existan
                    try:
                        os.remove(temp_file_path)
                    except FileNotFoundError:
                        print(f"El archivo {temp_file_path} no se encontró y no pudo ser eliminado.")

            # Enviar el archivo generado al usuario como descarga
            return send_file(documento_path, as_attachment=True)

    