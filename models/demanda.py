from docxtpl import DocxTemplate
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
import os
import tempfile
from flask import request, send_file
from werkzeug.utils import secure_filename
from models.calculos import formatear_dinero, transformar_fecha


class Formulario:
    def __init__(self, datos):
        """Inicializa la clase con el diccionario de datos del formulario."""
        self.datos = datos

    def datos_cliente(self):
        """Recoge los datos del cliente."""
        genero = ""
        if self.datos["datos_cliente"].get('genero') == "Masculino":
            genero = "del Sr"
        elif self.datos["datos_cliente"].get('genero') == "Femenino":
            genero = "de la Sra"
        return {
            'genero': genero,
            'nombre': self.datos["datos_cliente"].get('nombre'),
            'dni': self.datos["datos_cliente"].get('dni'),
            'domicilio': self.datos["datos_cliente"].get('domicilio'),
            'localidad': self.datos["datos_cliente"].get('localidad'),
            "fecha_adquisicion_derecho": transformar_fecha(self.datos["datos_cliente"].get("fecha_adquisicion_derecho")),
            "garciaVidal" : self.datos["datos_cliente"].get("garciaVidal")
        }

    def beneficio(self):
        """Recoge los datos de beneficios."""
        return {
            'fecha_reajuste': transformar_fecha(self.datos["beneficio"].get('fecha_reajuste')),
            'expediente_reajuste': self.datos["beneficio"].get('expediente_reajuste'),
            'numero_beneficio': self.datos["beneficio"].get('numero_beneficio'),
            'fecha_inicio_remuneraciones': transformar_fecha(self.datos["beneficio"].get('fecha_inicio_remuneraciones')),
            'fecha_fin_remuneraciones': transformar_fecha(self.datos["beneficio"].get('fecha_fin_remuneraciones')),
            'fecha_cese': transformar_fecha(self.datos["beneficio"].get('fecha_cese')),
            'ultima_remuneracion_actividad': formatear_dinero(self.datos["beneficio"].get('ultima_remuneracion_actividad')),
            'fecha_ultima_remuneracion_actividad': transformar_fecha(self.datos["beneficio"].get('fecha_ultima_remuneracion_actividad')),
            'ultima_remuneracion_actualizada_anses': formatear_dinero(self.datos["beneficio"].get('ultima_remuneracion_actualizada_anses')),
            'fecha_alta_primer_haber': transformar_fecha(self.datos["beneficio"].get('fecha_alta_primer_haber')),
            'monto_primer_haber': formatear_dinero(self.datos["beneficio"].get('monto_primer_haber')),
            'taza_de_reemplazo': self.datos["beneficio"].get('taza_de_reemplazo'),
            'ultimo_haber': formatear_dinero(self.datos["beneficio"].get('ultimo_haber')),
            'fecha_ultimo_haber': transformar_fecha(self.datos["beneficio"].get('fecha_ultimo_haber')),
            'fecha_reclamo' : transformar_fecha(self.datos["beneficio"].get('fecha_reclamo')),
        }

    def servicios(self):
        """Recoge los datos de servicios."""
        # Inicializar los textos de servicios
        servicios_dependencia = "Servicios en Dependencia: No tiene"
        servicios_autonomos = "Servicios Autonomos: No tiene"

        # Comprobar si se deben incluir servicios en dependencia
        if self.datos["servicios"].get('servicios_dependencia'):
            servicios_dependencia = "Cargo desempeñado y empleador al cese: " + \
                self.datos["servicios_dependencia"].get('cargo_desempleado', "") + " en " + \
                self.datos["servicios_dependencia"].get('empleador', "")

        # Comprobar si se deben incluir servicios autónomos
        if self.datos["servicios"].get('servicios_autonomos'):
            servicios_autonomos = "Servicios Autónomos: " + \
                self.datos["servicios_autonomos"].get('autonomo_input1', "") + " en " + \
                self.datos["servicios_autonomos"].get('autonomo_input2', "")

        # Devolver un diccionario con los resultados
        return {
            'servicios_dependencia': servicios_dependencia,
            'servicios_autonomos': servicios_autonomos,
        }

    def sumas_no_remunerativas(self):
        """Recoge los datos de servicios."""
        # Inicializar los textos de servicios
        Titulo_sumas_no_remunerativas = ""
        parrafo_introduccion = ""
        parrafo_sumas=""
        parrafo_legalidad1 = ""
        parrafo_legalidad2 = ""
        parrafo_legalidad3 = ""
        parrafo_legalidad4 = ""
        parrafo_legalidad5 = ""
        titulo_recibos = ""
        parrafo_recibos = ""
        parrafo_recibos_1 = ""
        genero = ""
        if self.datos["datos_cliente"].get('genero') == "Masculino":
            genero = "el Sr"
        elif self.datos["datos_cliente"].get('genero') == "Femenino":
            genero = "la Sra"

        if self.datos["casillas_verificacion"].get('opcion_sumas_remunerativas', False):
            # Título de la sección de sumas no remunerativas
            Titulo_sumas_no_remunerativas = "6.1.1.	De las sumas no remunerativas"

            parrafo_introduccion = (
                "Solicito se incorporen para el cálculo del ingreso base las sumas no remunerativas percibidas por mi mandante "
                "con carácter de normal y habitual por parte de su empleadora, provincia de Salta, conforme a la doctrina sentada "
                "en el caso 'Rainone de Ruffo' de la CSJN. Se acompaña a la presente la historia laboral de mi mandante, "
                "en donde se observa una columna que dice 'remuneración total', que es lo liquidado de mi mandante (incluye las "
                "sumas no remunerativas) y una que dice 'remuneración', que es sobre lo que aportó su empleador, ocasionándole "
                "un perjuicio a mi mandante."
            )

            # Párrafo de legalidad
            parrafo_legalidad1 = (
                "La Corte Suprema de Justicia de la Nación reconoció que el monto de las sumas no remunerativas debe "
                "ser considerado por ANSES para el cómputo del beneficio. Así, en la causa 'Rainone de Ruffo, Juana Teresa "
                "Berta c/ ANSeS s/ reajustes varios', Sentencia del 02.03.2011, donde se trataba de sumas no remunerativas "
                "abonadas por el propio ANSES como empleador, el Tribunal sostuvo que correspondía '(...) admitir la pretensión "
                "de la recurrente y ordenar que dichos montos sean incorporados en el cálculo del haber inicial ordenado por el "
                "juez de primera instancia, sin perjuicio del cargo por aportes omitidos y de las contribuciones que deban "
                "realizarse con destino a la seguridad social'. Máxime cuando el propio Organismo había reconocido como "
                "'remuneraciones sin aporte' las sumas en cuestión."
            )

            # Añadir el resto de los argumentos
            parrafo_legalidad2 = (
                " En consecuencia, al tratarse de sumas no remunerativas percibidas con carácter normal y habitual, corresponde "
                "considerar su monto para el cálculo de la jubilación, fundamentando que la misma ley de jubilaciones indica en "
                "su artículo 6 que 'a los fines previsionales, remuneración es todo ingreso que recibe un trabajador en retribución "
                "o compensación por su actividad personal prestados en relación de dependencia, incluidos los suplementos que "
                "tengan el carácter de habituales y regulares'."
            )

            parrafo_legalidad3 = (
                " Los pagos se hicieron con regularidad (variando el porcentaje respecto del salario). La omisión de aportes y "
                "contribuciones, por el eventual incumplimiento de los deberes a cargo de la Administración como agente de "
                "retención, no puede cambiar la verdadera naturaleza del desembolso efectuado. Además, la liberación de todo cargo "
                "al Estado Nacional en la implementación de los incentivos se refiere al origen de los fondos para llevar adelante "
                "el programa (arts. 4, 5 y 11 de la ley 23283), pero no lo libera de otras obligaciones, entre las que se encuentran "
                "las de obrar como agente de retención de los aportes y realizar las cotizaciones de seguridad social."
            )

            parrafo_legalidad4 = (
                " Más aún, el carácter remunerativo de los conceptos abonados encuadra en el art. 6 de la ley 24241, que asigna esa "
                "naturaleza 'a ciertas sumas que son abonadas a agentes de la Administración Pública, entre las que menciona al "
                "'premio estímulo, gratificaciones u otros conceptos de análogas características', con la modalidad de poner a cargo "
                "del agente, además de su aporte personal, la contribución que corresponde al empleador."
            )

            parrafo_legalidad5 = (
                " En virtud de lo expuesto, solicito se incorporen al cómputo del haber jubilatorio de mi mandante las sumas percibidas "
                "como no remunerativas y se ordene a su empleadora realizar las contribuciones previsionales correspondientes, teniendo "
                "en cuenta el criterio de ambas salas sobre este tema: Van Cauwlaert, Eduardo, sent. del 9/6/17, haciendo mérito de la "
                "doctrina que emana del precedente 'Rainone de Ruffo, Juana Teresa Berta' (Fallos: 334:210), y Fallos: 333:699 "
                "'González, Martín Nicolás c/ Polimat S.A. y otro', sent. del 19/5/2010."
            )

        # Comprobar si se deben incluir servicios en dependencia
        if self.datos["sumas_no_remunerativas"]["recibos"]["Recibos_Si"]:
            parrafo_sumas = "Se adjunta equiparación de haberes, de la que surge que el cargo desempeñado por mi representado al cese fue de " + self.datos["sumas_no_remunerativas"]["recibos_si"]["Cargo_Desempeñado"] + " en " + self.datos["sumas_no_remunerativas"]["recibos_si"]["Lugar_Desempeñado"] + ", con una antigüedad de " + self.datos["sumas_no_remunerativas"]["recibos_si"]["Años_antiguedad"] + " años de servicios.Asimismo, se adjuntan recibos de sueldo, de los que surgen que el empleador abonó a mi representado, haberes sin aportes bajo los siguientes códigos y conceptos, los que no fueron considerados para el cálculo del haber inicial"

        # Comprobar si se deben incluir servicios autónomos
        if self.datos["sumas_no_remunerativas"]["recibos"]["Recibos_No"]:
           parrafo_sumas = "Asimismo, peticiono se libre oficio a " + self.datos["sumas_no_remunerativas"]["recibos_no"]["Librar_oficio_a"] + " a los fines de que remita los recibos de sueldo de mi representada que se encuentran en su poder, correspondientes al período " + transformar_fecha(self.datos["sumas_no_remunerativas"]["recibos_no"]["inicio_periodo_sumas"]) + " hasta " + transformar_fecha(self.datos["sumas_no_remunerativas"]["recibos_no"]["fin_periodo_sumas"]) + " , de los que surgirán las sumas abonadas como no remunerativas, En su defecto, peticiono informe los haberes con aportes y sin aportes abonados en cada período peticionado.  De ellos surgirán las sumas no remunerativas abonadas por el empleador, bajo los siguientes códigos y conceptos: "
           titulo_recibos = "1.2 Prueba en poder de tercero:"
           parrafo_recibos = self.datos["sumas_no_remunerativas"]["recibos_no"]["Librar_oficio_a"]
           parrafo_recibos_1 = "Peticiono se libre oficio al empleador de mi representado,"+self.datos["sumas_no_remunerativas"]["recibos_no"]["Librar_oficio_a"]+", a los fines de que adjunte los recibos de sueldo correspondientes a "+genero+" "+self.datos["datos_cliente"].get('nombre')+" DNI "+self.datos["datos_cliente"].get('dni')+" por el período "+transformar_fecha(self.datos["sumas_no_remunerativas"]["recibos_no"]["inicio_periodo_sumas"])+" a "+transformar_fecha(self.datos["sumas_no_remunerativas"]["recibos_no"]["fin_periodo_sumas"])+", los que obran en su poder. En su defecto, peticiono informe los haberes con aportes y sin aportes abonados en todos los meses comprendidos en el período solicitado, los que influyen en el cálculo del haber inicial. Y lo sea bajo apercibimiento de astreintes en caso de incumplimiento."
        # Devolver un diccionario con los resultados
        return {
            'parrafo_sumas': parrafo_sumas,
            'parrafo_legalidad1': parrafo_legalidad1,
            'parrafo_legalidad2': parrafo_legalidad2,
            'parrafo_legalidad3': parrafo_legalidad3,
            'parrafo_legalidad4': parrafo_legalidad4,
            'parrafo_legalidad5': parrafo_legalidad5,
            'Titulo_sumas_no_remunerativas' : Titulo_sumas_no_remunerativas,
            'parrafo_introduccion': parrafo_introduccion,
            'titulo_recibos': titulo_recibos,
            'parrafo_recibos': parrafo_recibos,
            'parrafo_recibos_1': parrafo_recibos_1,
        }

    def error_material(self):
        """Recoge los datos de servicios."""
        # Inicializar los textos de servicios
        Titulo_error_material = ""
        parrafo_1 = ""
        parrafo_2 =""
        parrafo_3 = ""
        parrafo_4 = ""
        parrafo_5= ""
        parrafo_6= ""
        parrafo_7 = ""


        if self.datos["casillas_verificacion"].get('opcion_error_material', False):
            # Título de la sección de sumas no remunerativas
            Titulo_error_material= "5.1.2 Del error material "

            parrafo_1 = (
                "Mi mandante trabajó en " + self.datos["error_material"]["Lugar_error"] + " desde el " + transformar_fecha(self.datos["error_material"]["fecha_inicio_remuneraciones_error"]) + " hasta el " + transformar_fecha(self.datos["error_material"]["fecha_fin_remuneraciones_error"]) +"."
            )

            # Párrafo de legalidad
            parrafo_2 = (
                "Del detalle de beneficios de Anses se observan los siguientes errores materiales en los que incurrió el organismo previsional al momento del cálculo del haber jubilatorio inicial:"
            )

            # Añadir el resto de los argumentos
            parrafo_3= (
                "• Toma remuneraciones erróneas, diferentes a las efectivamente percibidas:"
            )

            parrafo_4 = (
                "• Se adjunta cálculo de haber de caja con y sin corrección del error material, de los que surgen los siguientes promedios de remuneraciones:"
            )

            parrafo_5 = (
                "o W de caja con error material en remuneraciones consideradas: " + formatear_dinero(self.datos["error_material"]["W_error"]) +"."
            )

            parrafo_6 = (
                "o W de caja sin error material, con remuneraciones correctas: " + formatear_dinero(self.datos["error_material"]["W_sin_error"]) +"."
            )
            parrafo_7 = (
                "Solicito se corrija el error material y se tomen las verdaderas remuneraciones percibidas para el cálculo del haber inicial." 
            )

        
        return {
            'Titulo_error_material': Titulo_error_material,
            'parrafo_1': parrafo_1,
            'parrafo_2': parrafo_2,
            'parrafo_3': parrafo_3,
            'parrafo_4': parrafo_4,
            'parrafo_5': parrafo_5,
            'parrafo_6': parrafo_6,
            'parrafo_7': parrafo_7,
        }

    def reajuste_pbu(self):
        """Recoge los datos de servicios."""
        # Inicializar los textos de servicios
        Titulo_PBU = ""
        parrafo_1 = ""
        parrafo_2 =""
        parrafo_3 = ""
        parrafo_4 = ""
        parrafo_5= ""
        parrafo_6= ""
        parrafo_7 = ""
        parrafo_8 = ""
        parrafo_9 = ""
        parrafo_10 = ""
        parrafo_11 = ""
        parrafo_12 = ""
        parrafo_13 = ""
        parrafo_14 = ""
        parrafo_15 = "" #aca se agregan cosas
        parrafo_16 = ""
        parrafo_17 = ""


        if self.datos["casillas_verificacion"].get('opcion_reajuste_pbu', False):
            Titulo_PBU= "5.1.4 PBU"

            parrafo_1 = (
                "Solicito se declare la inconstitucionalidad del art. 4° de la Ley 26.417 que estableció un monto fijo para la PBU y se utilice para la actualización de la Prestación Básica Universal, ISCIB al 02.2009. "
            )

            parrafo_2= (
                "Debe tenerse presente que al momento de la sanción de la Ley 24241 la determinación de esta prestación estaba ligada al valor del AMPO (arts. 20 y 21), que fue reemplazado por el MOPRE en el año 1997 (art. 1 Decreto 833/97) y alcanzó la suma de $80 a partir de abril de 1997 hasta septiembre de ese año, cuando dejó de publicarse y se mantuvo inalterado ($200) hasta la sanción de la Ley 26417 que elevó el monto de la PBU a $326."
            )

            # Añadir el resto de los argumentos
            parrafo_3 = (
              "Por lo tanto, el cálculo de la PBU determinado por ANSES surge palmariamente desactualizado, y por ello solicito la actualización del AMPO/MOPRE con el índice de salarios de la industria y la construcción (ISBIC), hasta el 28.02.2009, conforme el fallo “Aguado, Nélida del Carmen c/ ANSES y/o PEN s/ Reajustes Varios” Expte. N° FSA 15100230/2012, sentencia del 12.06.2019, “Fernández Gladis” FSA 18234/2014, sentencia del 19.06.2019 y “Jaureguina, Víctor Hugo” FSA 4900/2016, sentencia del 21.08.2019. Ello así, toda vez que es el mismo índice que se solicita para la actualización de las otras prestaciones (PC-PAP)."
            )

            parrafo_4 = (
                "Solicito tenga presente que desde la sanción de la ley N° 27.426, los índices de movilidad y actualización de remuneraciones fueron diferentes y se mantuvo el concepto de monto fijo. "
            )

            parrafo_5 = (
                "Luego, con la suspensión de la fórmula de movilidad jubilatoria, se adiciono un monto fijo en el haber para el mes de marzo de 2020, que luego en junio fue imputado en parte a la Pbu. "
            )

            parrafo_6 = (
                "En esencia, a la fecha del presente reclamo, la  PBU no guarda la proporción que tuvo en miras el legislador al crear dicho instituto, siendo determinación de la misma  regresiva y afectar la integralidad del haber de mi mandante, conforme lo acredito con las pruebas adjuntadas en autos."
            )
            parrafo_7 = (
                "Tenga presente VS. que la CSJN ha zanjado el tema respecto de que debe actualizarse a beneficios anteriores y posteriores a la sanción de la ley 26.417, poniendo especial énfasis en que todos los componentes del haber jubilatorio resultan revisables y que debe acreditarse la confiscatoriedad. "
            )
            parrafo_8 = (
                "Solicito que habiéndose  acreditado  el 15% de confiscatoriedad requerido en el fallo Conforme fallo “Quiroga” ( 337:1277) “Ciuti Pablo c/ ANSES s/ reajustes varios”, sentencia del 30/6/2015(CSJ 111/2012(48-C)/CS1);  Pichersky Alberto Raúl c/Anses s/reajustes Varios”, la C.S.J.N, el 23 de mayo de 2017(Expte SS 80278/20l2/l/RH 1)“González Héctor Orlando c/ ANSES s/ Reajuste de haberes” Expte FMP 41051103/2011/1/RH1., el reajuste del haber se haga contra pc y pap sin reajustar, dejando de lado los criterios fijados en Soule y Blanco , por cuanto cae la lógica de medición establecida en los mismos cuando la persona solo reclama el reajuste del PBU, a lo que se agrega que no se puede medir quita o merma respecto de una prestación ya mermada , conforme lo expresado por el Cuerpo de peritos de la CSJN, que determinó que el reajuste debe hacerse PBU reaj. + PC sin reaj. + PAP sin reaj. Solicito libre oficio a fin de solicitar a dicho organismo proceda a remitir copia de lo dictaminado sobre este punto."
            )
            parrafo_9 = (
                "De la lectura del precedente de la CSJN “Quiroga” se observa que la comparación debe hacerse con el total del haber inicial  y no con el haber total reajustado, conforme el considerando 10, que textualmente dice:"
            )
            parrafo_10 = (
                "LA CSJN no dice que el haber inicial sobre el que hay que medir deba contener PC y PAP reajustadas."
            )
            parrafo_11 = (
                "Además, no se puede medir quita o merma respecto de una prestación ya mermada, puesto que el haber inicial sobre el que hay que medir no puede contener PC y PAP reajustadas pues se estarían incorporando variables que no hacen a la naturaleza del componente sobre el que se quiere evaluar la incidencia de la quita."
            )
            parrafo_12 = (
                "Considerándose la medición el haber inicial -conforme considerando 10 del fallo- en la comparación entre haber de caja con PBU reajustada con ISBIC y PBU sin reajustar, solicito se reajuste la PBU sin realizar quita alguna, máxime teniendo en cuenta:"
            )
            parrafo_13 = (
                "a.	que en caso de que no se produzca una modificación en las remuneraciones al actualizarlas, cae la lógica fijada por esta jurisdicción en Soule y Blanco, por cuanto la metodología para evaluar la confiscatoriedad es compararla con PC y PAP reajustada "
            )
            parrafo_14 = (
                "b.	No se puede medir la merma con respecto a algo ya mermado (PBU sin reajustar).  Dicha medición pierde objetividad, sería como decir que la CSJN en “Del Azar Suaya” o “Actis Caporale” ordenó medir la incidencia de la quita sobre el tope y no sobre el haber reajustado."
            )
            parrafo_15 = (
                "c.	El haber sería más integral, teniendo en cuenta el último haber percibido por mi mandate en actividad que era de "+ formatear_dinero(self.datos["beneficio"].get('ultima_remuneracion_actividad')) + "  y la Pbu sin quita, permitiría obtener un haber de reemplazo del " + self.datos["PBU"].get('porcentaje_haber_reemplazo') + " y con quita de Soule del " + self.datos["PBU"].get('porcentaje_quita_Soule') + "."
            )
            parrafo_16 = (
                "d.	Las palabras quita, merma o disminución, ya tienen una quita, merma o disminución en el haber, por lo cual si quiero medir la incidencia tiene que ser sobre el haber antes de la disminución, por eso se debe tomar el haber de caja y no el reajustado."
            )
        if self.datos["PBU"]["quita_menor_15"]["quita_menor_15"]:
            parrafo_17 = "Por último, corresponde poner en resalto que si bien de acuerdo a los cálculos adjuntos la PBU no alcanza un 15% de confiscatoriedad, de no aplicarse su actualización arrojaría una quita del " + self.datos["PBU"]["quita_menor_15"]["quita_menor_15_1"] + ", por lo que de aplicarse los topes cuya inaplicabilidad de solicita, no deberían superar el "+ self.datos["PBU"]["quita_menor_15"]["quita_menor_15_2"] + "."

        return {
            'Titulo_PBU': Titulo_PBU,
            'parrafo_1': parrafo_1,
            'parrafo_2': parrafo_2,
            'parrafo_3': parrafo_3,
            'parrafo_4': parrafo_4,
            'parrafo_5': parrafo_5,
            'parrafo_6': parrafo_6,
            'parrafo_7': parrafo_7,
            'parrafo_8': parrafo_8,
            'parrafo_9': parrafo_9,
            'parrafo_10': parrafo_10,
            'parrafo_11': parrafo_11,
            'parrafo_12': parrafo_12,
            'parrafo_13': parrafo_13,
            'parrafo_14': parrafo_14,
            'parrafo_15': parrafo_15,
            'parrafo_16': parrafo_16,
            'parrafo_17': parrafo_17,
            
        }

    def taza_complementacion(self):
        """Recoge los datos de servicios."""
        # Inicializar los textos de servicios
        Titulo_taza_complementacion = ""
        parrafo_1 = ""
        parrafo_2 =""
        Titulo_parrafo_3 = ""
        parrafo_3 = ""
        Titulo_parrafo_4 = ""
        parrafo_4 = ""
        Titulo_parrafo_5 = ""
        parrafo_5 = ""
        parrafo_6 = ""
        parrafo_7 = ""
        parrafo_8 = ""

        if self.datos["casillas_verificacion"].get('opcion_tasa_complementacion', False):

            Titulo_taza_complementacion = "5.1.5 Tasa de complementación:  "
            
            parrafo_1 = "Atento a que, aun reajustando el haber de mi mandante conforme los parámetros ut supra solicitados, éste no guarda una debida proporcionalidad con el haber en actividad, solicito se fije una pauta de complementación que consagre la sustitutividad del haber de pasividad respecto al de actividad a fin de cumplir con la manda constitucional de tener una haber integral y en consecuencia sustitutivo del salario, en consonancia con los compromisos internacionales asumidos por el estado nacional al suscribir las normas mínimas de seguridad social de la OIT."

            parrafo_2 = "Los derechos a una retribución justa y a un salario mínimo vital y móvil encuentran su correlato en las jubilaciones y pensiones móviles que deben ser garantizadas a los trabajadores cuando entran en pasividad, nótese que el haber que percibe mi mandante es inferior al Salario mínimo vital y móvil"

            Titulo_parrafo_3 = "• SMVM "
            parrafo_3 = "A partir del 1º de julio de 2024 el salario mínimo es de $ 254.231,91 el cual ha decrecido notoriamente, no obstante, es superior a la jubilación mínima que percibe mi mandante. "

            if self.datos["Taza_complementacion"]["haber_menor_ripte"]:
                Titulo_parrafo_4 = "• Ripte"
                parrafo_4 = "al último índice aplicado, mayo de 2024 es de $879.483,08. "
                

            Titulo_parrafo_5 = "• Canasta Básica Total "            
            parrafo_5 = "$ 282.579 a Julio de 2024."

            parrafo_6 = "La Cámara Federal de Salta también se pronunció en un fallo de la Sala II, en autos “GOMEZ AUGIER GUSTAVO FEDERICO C/ ANSES S/ REAJUSTES VARIOS” EXPTE. 11730/2016, donde, haciendo alusión a los casos “Betancur” y “Benoist”, sentó doctrina respecto a la integralidad de los derechos de la seguridad social reconocido en el art. 14 bis de la CN, y sostuvo que “(…) si bien no corresponde fijación de una “tasa” de sustitución para que el beneficio de jubilación ordinaria otorgado al actor bajo el régimen de la ley 24.241 alcance un mínimo determinado - tal como lo establecía el art. 49 de la ley 18.037-, ello no enerva el derecho del accionante de acreditar en la etapa de ejecución la necesidad de establecer un suplemento que resguarde los principios de “sustitutividad” y de “proporcionalidad” que, según los lineamientos del Superior Tribunal, debe existir entre la jubilación y el ingreso que tenía cuando se encontraba en actividad” ."

            parrafo_7 = "En este sentido, el Tribunal de Alzada de nuestra jurisdicción entendió que “… si luego de la redeterminación del haber de inicio conforme las pautas de sentencia y efectuada la verificación de confiscatoriedad -tanto de la merma producida ante la ausencia de incrementos de la Prestación Básica Universal, como de la aplicación de los topes máximos-el análisis integral del haber reajustado demuestra que el haber de pasividad no guarda razonable proporción con el haber de actividad ejercido al cese por el titular, corresponderá establecer -como última ratio- una pauta de complementación del beneficio que torne operativa la directriz jurídica no normativa que dimana de los principios de sustitutividad y proporcionalidad”."

            parrafo_9 = "Por lo que así lo solicito, a fin de que el haber de mi mandante sea integral y sustitutivo del salario."

            if self.datos["Taza_complementacion"]["print_equiparacion"]["print_equiparacion"]:
                calculo = self.datos["Taza_complementacion"]["print_equiparacion"]["monto_jubilacion_taza"] * 100
                calculo = calculo / self.datos["Taza_complementacion"]["print_equiparacion"]["monto_cargo_ejercido"]
                calculo = round(calculo, 2)
                parrafo_8 = "Note VS. que surge de la equiparación expedida por su ex empleador, " + self.datos["Taza_complementacion"]["print_equiparacion"]["taza_ex_empleador"] + ", que siendo su última categoría de " + self.datos["Taza_complementacion"]["print_equiparacion"]["ultimo_cargo_ejercido"] + " , por el mes " + self.datos["Taza_complementacion"]["print_equiparacion"]["fecha_monto_cargo_ejercido"] + " su remuneración hubiese sido por el monto de " + formatear_dinero(self.datos["Taza_complementacion"]["print_equiparacion"]["monto_cargo_ejercido"]) +", mientras que el organismo previsional determinó en el mismo mes un haber jubilatorio de "+ formatear_dinero(self.datos["Taza_complementacion"]["print_equiparacion"]["monto_jubilacion_taza"]) +", el que solo representa un "+str(calculo)+"% de lo que le correspondería de estar en actividad, por lo tanto, no es sustitutivo del salario activo y no alcanza la tasa de complementación que se peticiona "

            return {
                "Titulo_taza_complementacion": Titulo_taza_complementacion,
                "parrafo_1": parrafo_1,
                "parrafo_2": parrafo_2,
                "Titulo_parrafo_3": Titulo_parrafo_3,
                "parrafo_3": parrafo_3,
                "Titulo_parrafo_4": Titulo_parrafo_4,
                "parrafo_4": parrafo_4,
                "Titulo_parrafo_5": Titulo_parrafo_5,
                "parrafo_5": parrafo_5,
                "parrafo_6": parrafo_6,
                "parrafo_7": parrafo_7,
                "parrafo_8": parrafo_8,
                "parrafo_9": parrafo_9
            }
    def Inaplicabilidad_tope_Art_14_de_la_Res_06_09(self):
        titulo_inaplicabilidad_tope_art_14_de_la_res_06_09 =""
        parrafo_1=""
        if self.datos["casillas_verificacion"].get('opcion_inaplicabilidad_tope_art_14_res_06_09', False):
            titulo_inaplicabilidad_tope_art_14_de_la_res_06_09 = "5.1.6. Inaplicabilidad tope Art 14 de la Res 06.09"
            parrafo_1 = "Solicito su inaplicabilidad por la cual al  establecer un tope a las actualizaciones de las remuneraciones, excede la facultad reglamentaria- ya que el propio artículo 24 de la ley 24.241 no establece tope alguno-, no debería incidir en este caso , dado que el mismo se fijó para actualizar las remuneraciones que componen la PC -no la Pap-, y conforme el periodo usado para el promedio de remuneraciones, no se aplica, y no es justo que se limite el haber sobre el cual se aportó, dado que arrojaría un resultado disvalioso."

        return {
            'titulo' : titulo_inaplicabilidad_tope_art_14_de_la_res_06_09,
            'parrafo_1' : parrafo_1
        }
    def Inaplicabilidad_del_tope_del_art_9_y_25_de_la_ley_24241(self):
        titulo_inaplicabilidad_del_tope_del_art_9_y_25_de_la_ley_24241 =""
        parrafo_1=""
        parrafo_2_1=""
        parrafo_2_2=""
        parrafo_3=""
        parrafo_4=""
        genero = ""
        if self.datos["datos_cliente"].get('genero') == "Masculino":
            genero = "el Sr"
        elif self.datos["datos_cliente"].get('genero') == "Femenino":
            genero = "la Sra"
        if self.datos["casillas_verificacion"].get('opcion_inaplicabilidad_tope_ley_24241', False):
            titulo_inaplicabilidad_del_tope_del_art_9_y_25_de_la_ley_24241 = "5.1.7 Inaplicabilidad del tope del art. 9 y 25 de la ley 24.241"
            parrafo_1 = "Solicito se declare la inconstitucionalidad e inaplicabilidad del tope del art. 9 y 25 de la ley 24.241, por cuanto su aplicación para el cómputo de la PC y PAP implica la no consideración de los importes superiores a la remuneración máxima imponible, lo que disminuye notoriamente su nivel de ingresos a la hora de determinar la prestación previsional sustitutiva del haber en actividad."
            parrafo_2_1 = "Peticiono se aplique lo recientemente resuelto por la Cámara Federal de Salta – Sala II, en los autos caratulados “RODRIGUEZ, MARIA CRISTINA c/ANSES s/REAJUSTE DE HABERES” Expte. N° FSA 13956/2015, sentencia de fecha 23 de agosto de 2024, la que entre sus considerandos sostuvo primeramente sobre la base de lo resuelto por el Máximo tribunal en el fallo “Gualtieri Alberto” (Fallos: 340:411): "
            parrafo_2_2 = "“debemos tener en cuenta que a los fines del cálculo de la prestación jubilatoria de la Sra. Rodríguez se computaron remuneraciones por el período comprendido entre septiembre de 2003 y agosto de 2013, no surgiendo de las constancias de la causa que la actora, encontrándose en actividad y vigente el anterior régimen de capitalización, solicitara aportar por un importe superior al límite del art. 9 de la ley 24.241. Sin embargo, no puede achacársele igual conducta desaprensiva de su futuro previsional con respecto al período posterior a noviembre de 2008, considerando que el anterior régimen de capitalización estuvo vigente hasta el 8 de diciembre de 2008, ya que a partir de allí la accionante no contaba con la posibilidad legal de efectuar imposiciones voluntarias o depósitos convenidos.”"
            parrafo_3="Así, seguidamente resolvió declarar la inconstitucionalidad de los arts. 9 y 25 de la ley 24.241, en el período comprendido desde diciembre de 2008 en adelante, para el re-cálculo del haber inicial. "
            parrafo_4="Teniendo en cuenta que en el presente caso para el cálculo del haber inicial se consideraron las remuneraciones percibidas por "+ genero + " "+ self.datos["datos_cliente"].get('nombre') + " desde el "+transformar_fecha(self.datos["beneficio"].get('fecha_inicio_remuneraciones'))+" a " + transformar_fecha(self.datos["beneficio"].get('fecha_fin_remuneraciones')) + ", es decir que todas sus remuneraciones consideradas para el cálculo de la PC y PAP son posteriores a diciembre de 2008, peticiono se declare la inconstitucionalidad y por lo tanto, la inaplicabilidad del tope previsto por los arts. 9 y 25 de la ley 24.241 por cuanto mi mandante no tuvo posibilidad alguna de ingresar aportes, facultad que solo estaba disponible en vigencia de la ley 24241 sin la modificación del régimen de capitalización."
        return {
            'titulo' : titulo_inaplicabilidad_del_tope_del_art_9_y_25_de_la_ley_24241,
            'parrafo_1' : parrafo_1,
            'parrafo_2_1' : parrafo_2_1,
            'parrafo_2_2' : parrafo_2_2,
            'parrafo_3' : parrafo_3,
            'parrafo_4' : parrafo_4
        }

    def tope_haber_maximo(self):
        titulo_tope_haber_maximo =""
        parrafo_1=""
        parrafo_2=""
        parrafo_3=""
        parrafo_4=""
        parrafo_5=""
        parrafo_6=""
        parrafo_7=""
        parrafo_8=""
        parrafo_9=""
        parrafo_10=""
        parrafo_11=""
        parrafo_12=""
        parrafo_13=""
        parrafo_14=""
        parrafo_15=""
        parrafo_16=""
        parrafo_17=""

        if self.datos["casillas_verificacion"].get('opcion_movilidad_tope_haber_maximo', False):
            titulo_tope_haber_maximo = "5.5 Tope de haber máximo "
            parrafo_1 = "En el caso, de que del resultado del juicio surja que el haber reajustado supera el tope establecido por ley otorgante, desde ya, dejó planteada su inconstitucionalidad, en concordancia con la doctrina de la Corte Suprema de Justicia de la Nación; en la medida que su aplicación implique una merma en el haber previsional que, por su magnitud, sea confiscatorio (cfr. sent. del 25.09.97, Del Azar Suaya, Abraham”). De allí que sólo procederá la tacha de inconstitucionalidad cuando se demuestre que la aplicación del tope legal importe una disminución irrazonable del haber de pasividad en relación al nivel de vida del titular, medido en función de la pauta legal contemplada en la ley mediante la cual se obtuvo el beneficio. "
            parrafo_2 = "El artículo 9 de la Ley 24.463 establece los llamados topes máximos, los que justifica en el principio de solidaridad que rige el sistema. "
            parrafo_3="Esta limitación a la percepción del haber, resulta lesiva al Art. 17 de la CN. En el supuesto de que la correspondiente liquidación arroje un haber que supere el tope reajustado, teniendo dicha limitación carácter confiscatorio, solicito expresamente: "
            parrafo_4="a)	Se aplique el criterio citado en Actis Caporale, Tudor y Pellegrini Américo, por cuanto si el tope excede el 15% no se debe aplicar."
            parrafo_5="b)	En caso de que decida aplicarse el mismo, expresamente se deje constancia que, de las sucesivas acumulaciones de tope, no puede producirse una quita superior al 15%, como se resolvió en “García Vidal”."
            parrafo_6="c)	 Previo a la aplicación del tope solicito que previamente se movilice el mismo conforme “Badaro”, por cuanto los $3.100 estuvieron fijos en el periodo 2002 a 2006, teniendo en cuenta que en su origen el mismo era el 82% de la remuneración máxima sujeta a aportes, y hoy con suerte llega al 60%. Debiendo aplicarse para movilizar este tope del art 9 inc. 3, el mismo criterio fijado por la CSJN en los autos “Badaro” para el periodo 2002 a 2006 y las pautas de movilidad que se fijen en la sentencia."
            parrafo_7="d)	En el improbable caso que no haga lugar al punto anterior, expresamente solicito que la movilidad que le fije a los haberes, se la den al tope del haber máximo."
            parrafo_8="El tope del art 9 de la ley 24463 es a agosto de 2024 de $1.517.094,79 (Res. ANSES 390.2024)."
            parrafo_9="Si se tuviera en cuenta que los $3.100 representaban el 82% de la remuneración máxima sujeta a aportes que era $3.780 (60 ampo de $63), y trajéramos el mismo criterio a hoy, nos daría que el tope del haber máximo debiera ser a agosto de 2024 del 82% de la remuneración máxima sujeta a aportes que es $2.467.787,04= $2.023.426,54"
            parrafo_10="Por lo que, si a los $3.100, que eran a 01/2002, lo actualizamos con las mismas pautas de “Badaro” y aumentos generales de ANSES el tope hoy sería notablemente superior."
            parrafo_11="•	Haber reclamado al "+ transformar_fecha(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["fecha_haber_actual"]) +" reajustado cf. Ley 27.551 (50% IPC y 50% RIPTE mensual) sin topes: " + formatear_dinero(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["reajustado_cf_ley_27551_mensual"])
            parrafo_12="•	Haber reclamado al "+ transformar_fecha(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["fecha_haber_actual"]) +" reajustado cf. Ley 27.551 (50% IPC y 50% RIPTE retrasado tres meses) sin topes: " + formatear_dinero(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["reajustado_cf_ley_27551_tres_meses"])
            parrafo_13="•	Haber reclamado al "+ transformar_fecha(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["fecha_haber_actual"]) +" reajustado cf. IPC retrasado dos meses sin topes: " + formatear_dinero(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["reajustado_IPC_sin_topes"])
            parrafo_14="•	Tope de haber máximo de ANSES sin actualizar al "+ transformar_fecha(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["fecha_haber_actual"]) +" : "+ formatear_dinero(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["tope_haber_maximo_anses"])
            parrafo_15="•	Tope actualizado cf. Badaro más Caliva Márquez: "+ formatear_dinero(self.datos["Tope_haber_maximo"]["Tope_haber_maximo"]["tope_actualizado_cf_badaro_mas_caliva_marquez"])
            parrafo_16="De aplicarse la movilidad del tope que se peticiona, ello permitiría mantener la integralidad del haber de mi mandante."
            parrafo_17="Acá se cuestiona el monto del haber máximo, como así también la quita que se produce tanto por su inadecuada movilidad y como por su aplicación, provocando un daño notorio en el haber, por cuanto si tengo que realizar una quita del 15% sobre un haber máximo no actualizado desde 1995 que se dictó la ley 24463, a mayo de 2006 que se incrementó con el Decreto Nº 764/06, un 11%, resultando el mismo de $3.441. Los aumentos posteriores otorgados a las prestaciones previsionales mediante la Ley 26.198, los Decretos 1346/07, 279/08 y posteriormente en virtud de la ley de movilidad 26.417, se reflejaron en el haber máximo. Pero de 2002 al 05-2006 NO, causando un perjuicio enorme a mi mandante en caso de ser alcanzado por el mismo."

        return {
            'titulo' : titulo_tope_haber_maximo,
            'parrafo_1' : parrafo_1,
            'parrafo_2' : parrafo_2,
            'parrafo_3' : parrafo_3,
            'parrafo_4' : parrafo_4,
            'parrafo_5' : parrafo_5,
            'parrafo_6' : parrafo_6,
            'parrafo_7' : parrafo_7,
            'parrafo_8' : parrafo_8,
            'parrafo_9' : parrafo_9,
            'parrafo_10' : parrafo_10,
            'parrafo_11' : parrafo_11,
            'parrafo_12' : parrafo_12,
            'parrafo_13' : parrafo_13,
            'parrafo_14' : parrafo_14,
            'parrafo_15' : parrafo_15,
            'parrafo_16' : parrafo_16,
            'parrafo_17' : parrafo_17
        }

    def ganancias(self):
        titulo_ganancias =""
        if self.datos["casillas_verificacion"].get('opcion_inaplicabilidad_impuesto_ganancias', False):
            titulo_ganancias="12.Ganancias: "
        return{
            'titulo' : titulo_ganancias
        }
        
    def generar_diccionario_docx(self):
        """Genera el diccionario que se usará para crear el archivo Word."""
        doc_data = {
            'cliente': self.datos_cliente(),
            'beneficio': self.beneficio(),
            'servicios': self.servicios(),
            'sumas_no_remunerativas': self.sumas_no_remunerativas(),
            'error_material': self.error_material()
            # Agrega más secciones según lo necesites
        }
        return doc_data

    def crear_archivo_word(self, imagenes_por_marcador):
        """Crea un archivo Word usando la plantilla y los datos proporcionados, incluyendo imágenes."""

        fecha_escrito = self.datos["datos_cliente"].get("fecha_adquisicion_derecho")
        fecha_escrito = datetime.strptime(fecha_escrito, '%Y-%m-%d')

        # Restar un año si garcia_vidal es True
        if self.datos["datos_cliente"].get("garciaVidal"):
            fecha_escrito = fecha_escrito - timedelta(days=365)

        # Seleccionar plantilla según la fecha_escrito
        if fecha_escrito < datetime(2018, 3, 1):
            plantilla = 'datos/MODELO ANT 03.2018. COMPLETO..docx'
        elif datetime(2018, 3, 1) <= fecha_escrito < datetime(2021, 1, 1):
            plantilla = 'datos/MODELO POST 3.18 COMPLETO..docx'
        else:
            plantilla = 'datos/MODELO POST 01.2021 COMPLETO.docx'

        # Cargar la plantilla y crear el contexto
        doc = DocxTemplate(plantilla)
        contexto = {
            'datos_cliente': self.datos_cliente(),
            'servicios': self.servicios(),
            'beneficio': self.beneficio(),
            'sumas_no_remunerativas': self.sumas_no_remunerativas(),
            'error_material': self.error_material(),
            'PBU': self.reajuste_pbu(),
            'taza_complementacion' : self.taza_complementacion(),
            'inaplicabilidad_tope_art_14_de_la_res_06_09' : self.Inaplicabilidad_tope_Art_14_de_la_Res_06_09(),
            'Inaplicabilidad_del_tope_del_art_9_y_25_de_la_ley_24241' : self.Inaplicabilidad_del_tope_del_art_9_y_25_de_la_ley_24241(),
            'tope_haber_maximo' : self.tope_haber_maximo(),
            'ganancias' : self.ganancias(),
        }

        # Renderizar el documento con el contexto
        doc.render(contexto)

        # Guardar el documento editado temporalmente
        temp_doc_path = 'datos/documento_temporal.docx'
        doc.save(temp_doc_path)

        # Manejo de las imágenes
        doc = Document(temp_doc_path)

        # Insertar imágenes en posiciones específicas según el diccionario
        for marcador, imagen_path in imagenes_por_marcador.items():
            for paragraph in doc.paragraphs:
                if marcador in paragraph.text:  # Buscar el marcador
                    paragraph.text = paragraph.text.replace(marcador, '')  # Eliminar el marcador del texto

                    if imagen_path:  # Verificar si hay imagen para ese marcador
                        run = paragraph.add_run()
                        run.add_picture(imagen_path, width=Inches(5))  # Ajusta el tamaño según sea necesario
                    break  # Detener la búsqueda en los párrafos una vez encontrado el marcador

        # Guardar el documento final editado
        final_path = 'datos/documento_editado.docx'
        doc.save(final_path)

        # Devolver la ruta del archivo final
        return final_path

    def procesar_imagenes(self):
        """Procesa las imágenes subidas por el usuario y las asigna a sus respectivos marcadores."""
        # Obtener las imágenes proporcionadas por el usuario (pueden ser opcionales)
        #sumas no remunerativas
        imagenSumas = self.datos["sumas_no_remunerativas"]["Imagen"].get("Imagen")

        # error material
        imagenError = self.datos["error_material"]["Imagen"].get("Imagen")

        # PBU
        imagenPBU_1 = self.datos["PBU"]["Imagen"].get("Imagen1")
        imagenPBU_2= self.datos["PBU"]["Imagen"].get("Imagen2")
        imagen_fija_pbu = 'datos/imagenes_fijas/imagen_fija_pbu.png'

        #Taza de complementacion
        Imagen_Taza = self.datos["Taza_complementacion"]["Imagen"].get("Imagen")
        imagen_fija_taza = 'datos/imagenes_fijas/imagen_fija_taza.png'

        #Tope Haber Maximo
        imagenTopeMaximo = self.datos["Tope_haber_maximo"]["Imagen"].get("Imagen")

        # Crear diccionario para asignar las imágenes a sus marcadores
        imagenes_por_marcador = {}

        # Asignar cada imagen a su respectivo marcador si existe
        for imagen, marcador in zip([imagenSumas, imagenError, imagenPBU_1, imagenPBU_2, Imagen_Taza,imagenTopeMaximo  ], ['Imagen_suma_aqui', 'Imagen_error_material_aqui', 'Imagen_PBU_1', 'Imagen_PBU_2', 'Imagen_Taza', 'Imagen_Tope_haber_maximo' ]):
            if imagen:  # Verificar si el usuario subió una imagen
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                    temp_file.write(imagen.read())
                    imagenes_por_marcador[marcador] = temp_file.name  # Mapea el marcador a la ruta del archivo temporal
            else:
                imagenes_por_marcador[marcador] = None  # No hay imagen para este marcador, se eliminará del documento

        #Imagen PBU
        if self.datos["casillas_verificacion"].get('opcion_reajuste_pbu', False):
            with open(imagen_fija_pbu, 'rb') as imagen_fija:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file_fija:
                    temp_file_fija.write(imagen_fija.read())
                    imagenes_por_marcador['Imagen_fija_pbu_aqui'] = temp_file_fija.name  # Agregar la imagen
        else:
            # Eliminar marcador 'Imagen_fija_pbu_aqui' si no hay imagen
            imagenes_por_marcador['Imagen_fija_pbu_aqui'] = ''  # Asignar cadena vacía si no hay imagen

        # Verificar si existe la imagen para el marcador 'Imagen_fija_taza_aqui'
        if self.datos["casillas_verificacion"].get('opcion_tasa_complementacion', False):
            with open(imagen_fija_taza, 'rb') as imagen_fija:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file_fija:
                    temp_file_fija.write(imagen_fija.read())
                    imagenes_por_marcador['Imagen_fija_taza_aqui'] = temp_file_fija.name  # Agregar la imagen
        else:
            # Eliminar marcador 'Imagen_fija_taza_aqui' si no hay imagen
            imagenes_por_marcador['Imagen_fija_taza_aqui'] = ''  # Asignar cadena vacía si no hay imagen

        # Llama a la función para crear el documento Word
        documento_path = self.crear_archivo_word(imagenes_por_marcador)  # Pasa el diccionario

        # Eliminar los archivos temporales después de usarlos
        for temp_file_path in imagenes_por_marcador.values():
            if temp_file_path:  # Solo eliminar archivos que existan
                try:
                    os.remove(temp_file_path)
                except FileNotFoundError:
                    print(f"El archivo {temp_file_path} no se encontró y no pudo ser eliminado.")

        # Enviar el archivo generado al usuario como descarga
        return send_file(documento_path, as_attachment=True)